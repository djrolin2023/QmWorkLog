# -*- coding: utf-8 -*-
"""乾明工作台账系统
默认端口 8088，默认账号/密码：admin / admin
数据目录: Data/[类型]/[年份]/[月份]/N.[类型]YYYY-MM-DD.docx
归档目录: [类型]/[年份]/[月份]（月度）、[类型]/[年份]（年度），可打包下载
"""
import io
import os
import re
import json
import hmac
import shutil
import hashlib
import zipfile
import functools
import collections
import calendar
from datetime import datetime, timedelta
from flask import (Flask, request, session, redirect, url_for, render_template,
                   jsonify, send_file, abort)

import docx_utils
from docx_utils import daily_filename
from docx import Document
from docx.shared import Pt, RGBColor
import config
import db
import logging
from paths import EXE_DIR, RES_DIR

BASE = EXE_DIR
DATA_DIR = os.path.join(BASE, 'Data')
USERS_FILE = os.path.join(BASE, 'users.json')
STATIC_IMAGES = os.path.join(RES_DIR, 'static', 'Images')
FNAME_RE = re.compile(r'^(\d{1,2})\.(.+?)(\d{4})-(\d{2})-(\d{2})\.docx$')
VERSION = '26.07.26'

# 应用级日志（写 app.log，便于排查，不影响主流程）
logging.basicConfig(
    filename=os.path.join(BASE, 'app.log'),
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    encoding='utf-8',
)
app_logger = logging.getLogger('qm_worklog')


def load_version():
    """读取 version.json 的当前版本号与历史记录。"""
    try:
        with open(os.path.join(RES_DIR, 'version.json'), 'r', encoding='utf-8') as f:
            d = json.load(f)
        return d.get('version', VERSION), d.get('history', [])
    except Exception:
        return VERSION, []


app = Flask(__name__,
            template_folder=os.path.join(RES_DIR, 'templates'),
            static_folder=os.path.join(RES_DIR, 'static'))
app.secret_key = config.load_secret_key()
app.config['MAX_CONTENT_LENGTH'] = 200 * 1024 * 1024
app.config['JSON_AS_ASCII'] = False
# 会话持久化：登录后 cookie 有效期 30 天，重启进程/关闭浏览器后仍保持登录。
# secret_key 为固定常量，重启后已签名的 cookie 继续有效。
app.config['SESSION_COOKIE_NAME'] = 'qmworklog_session'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=30)
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'


@app.context_processor
def _inject_version():
    ver, _ = load_version()
    return {'app_version': ver}


# 禁用浏览器缓存，避免修改 css/js/模板后页面仍显示旧版本
@app.after_request
def _no_cache(resp):
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    resp.headers['Pragma'] = 'no-cache'
    resp.headers['Expires'] = '0'
    return resp


# ------------------------- 用户 -------------------------

def _hash(pw):
    return hashlib.sha256(('qm#' + pw).encode('utf-8')).hexdigest()


def verify_password(pw, stored):
    return hmac.compare_digest(stored, _hash(pw))


def load_users():
    if not os.path.exists(USERS_FILE):
        users = {'admin': _hash('admin')}
        with open(USERS_FILE, 'w', encoding='utf-8') as f:
            json.dump(users, f)
        return users
    with open(USERS_FILE, 'r', encoding='utf-8') as f:
        return json.load(f)


def save_users(users):
    _atomic_write(USERS_FILE, lambda f: json.dump(users, f))


def parse_int(val, default=0):
    """容错地把字符串转 int，失败返回默认值（避免非数字参数导致 500）。"""
    if val is None:
        return default
    try:
        return int(str(val).strip())
    except (ValueError, TypeError):
        return default


def _atomic_write(path, writer):
    """原子写：先写 .tmp 临时文件，再 os.replace 替换，避免崩溃损坏文件。"""
    tmp = path + '.tmp'
    with open(tmp, 'w', encoding='utf-8') as f:
        writer(f)
    os.replace(tmp, path)


def list_user_names():
    return list(load_users().keys())


def add_user(username, password):
    users = load_users()
    if not username:
        return False, '用户名不能为空'
    if username in users:
        return False, '用户名已存在'
    users[username] = _hash(password)
    save_users(users)
    return True, '已添加'


def delete_user(username):
    users = load_users()
    if username not in users:
        return False, '用户不存在'
    if username == 'admin':
        return False, '默认管理员 admin 不可删除'
    if username == session.get('user'):
        return False, '不能删除当前登录账号'
    del users[username]
    save_users(users)
    return True, '已删除'


def change_password(username, old_pw, new_pw):
    users = load_users()
    if username not in users:
        return False, '用户不存在'
    if not verify_password(old_pw, users[username]):
        return False, '原密码错误'
    if not new_pw:
        return False, '新密码不能为空'
    users[username] = _hash(new_pw)
    save_users(users)
    return True, '密码已修改'


def reset_password(username, new_pw):
    users = load_users()
    if username not in users:
        return False, '用户不存在'
    if not new_pw:
        return False, '新密码不能为空'
    users[username] = _hash(new_pw)
    save_users(users)
    return True, '密码已重置'


def login_required(fn):
    @functools.wraps(fn)
    def wrapper(*a, **kw):
        if not session.get('user'):
            if request.path.startswith('/api/'):
                return jsonify(ok=False, msg='未登录'), 401
            return redirect(url_for('login'))
        return fn(*a, **kw)
    return wrapper


# ------------------------- 工具 -------------------------

def check_type(t):
    if not t or db.type_id(t) is None:
        abort(400, '类型错误')
    return t


def rel_path(path):
    return os.path.relpath(path, BASE).replace('\\', '/')


def sync_record(dtype, y, m, d, path, content=''):
    tid = db.type_id(dtype)
    if tid is not None:
        db.upsert_record(tid, f'{y:04d}-{m:02d}-{d:02d}',
                         rel_path(path), content)


def parse_date(s):
    try:
        dt = datetime.strptime(s, '%Y-%m-%d')
        return dt.year, dt.month, dt.day
    except Exception:
        abort(400, '日期格式错误')


def data_path(dtype, y, m, d):
    return os.path.join(DATA_DIR, dtype, f'{y:04d}', f'{m:02d}',
                        daily_filename(dtype, y, m, d))


def find_record(dtype, y, m, d):
    """按日期查找文件（兼容序号不同的命名）"""
    p = data_path(dtype, y, m, d)
    if os.path.exists(p):
        return p
    folder = os.path.dirname(p)
    if os.path.isdir(folder):
        suffix = f'{dtype}{y:04d}-{m:02d}-{d:02d}.docx'
        for f in os.listdir(folder):
            if f.endswith(suffix):
                return os.path.join(folder, f)
    return None


def zip_folder_to_memory(folder, arc_root):
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for root, _dirs, files in os.walk(folder):
            for f in files:
                full = os.path.join(root, f)
                rel = os.path.relpath(full, folder)
                zf.write(full, os.path.join(arc_root, rel))
    buf.seek(0)
    return buf


# ------------------------- 页面 -------------------------

@app.route('/login', methods=['GET', 'POST'])
def login():
    cfg = config.load_config()
    if request.method == 'POST':
        username = (request.form.get('username') or '').strip()
        password = request.form.get('password') or ''
        if not username or not password:
            return render_template('login.html', error='请输入用户名和密码', cfg=cfg)
        users = load_users()
        if username not in users:
            return render_template('login.html', error='用户名不存在', cfg=cfg)
        if not verify_password(password, users[username]):
            return render_template('login.html', error='密码错误', cfg=cfg)
        session['user'] = username
        session.permanent = True
        db.log_action(username, '登录', '系统', '登录成功', request.remote_addr)
        return redirect(url_for('index'))
    return render_template('login.html', cfg=cfg)


@app.route('/logout')
def logout():
    u = session.get('user')
    if u:
        db.log_action(u, '登出', '系统', '退出登录', request.remote_addr)
    session.clear()
    return redirect(url_for('login'))


def _render(page):
    return render_template(page, user=session['user'], cfg=config.load_config())


@app.route('/')
@login_required
def index():
    return redirect(url_for('dashboard'))


@app.route('/dashboard')
@login_required
def dashboard():
    # 打开数据大屏时自动增量同步索引（原"重建索引"按钮已移除，改为自动执行）
    try:
        db.sync_filesystem()
    except Exception as e:
        app.logger.warning('自动同步索引失败: %s', e)
    summary = db.stats_summary()
    by_type = db.stats_by_type()
    monthly = db.stats_monthly()
    recent = db.recent_records(8)
    total = summary['total']
    month_new = summary['month_new']
    other = total - month_new
    # 环形图：本月录入 / 历史录入
    donut_c = 2 * 3.141592653589793 * 50
    month_len = round(month_new / total * donut_c, 1) if total else 0
    other_len = round(other / total * donut_c, 1) if total else 0
    type_max = max((t['count'] for t in by_type), default=1) or 1
    month_max = max((m['count'] for m in monthly), default=1) or 1
    return render_template(
        'dashboard.html', user=session['user'], cfg=config.load_config(),
        summary=summary, by_type=by_type, monthly=monthly, recent=recent,
        month_new=month_new, other=other,
        donut_c=donut_c, month_len=month_len, other_len=other_len,
        type_max=type_max, month_max=month_max)


@app.route('/ledger')
@login_required
def ledger():
    return _render('ledger.html')


@app.route('/settings')
@login_required
def settings():
    return _render('settings.html')


@app.route('/users')
@login_required
def users():
    return _render('users.html')


@app.route('/types')
@login_required
def types():
    return _render('types.html')


@app.route('/about')
@login_required
def about():
    return _render('about.html')


@app.route('/manual')
@login_required
def manual():
    return _render('manual.html')


@app.route('/archive')
@login_required
def archive():
    return _render('archive.html')


@app.route('/impexp')
@login_required
def impexp():
    return _render('impexp.html')


@app.route('/summary')
@login_required
def summary():
    return _render('summary.html')


# ------------------------- 台账 API -------------------------

@app.route('/api/records')
@login_required
def api_records():
    """某类型某月已有记录的日期列表"""
    dtype = check_type(request.args.get('type', ''))
    y = parse_int(request.args.get('year', 0))
    m = parse_int(request.args.get('month', 0))
    folder = os.path.join(DATA_DIR, dtype, f'{y:04d}', f'{m:02d}')
    days = []
    if os.path.isdir(folder):
        for f in sorted(os.listdir(folder)):
            mt = FNAME_RE.match(f)
            if mt:
                days.append({'day': int(mt.group(5)), 'file': f,
                             'date': f'{mt.group(3)}-{mt.group(4)}-{mt.group(5)}'})
    days.sort(key=lambda x: x['day'])
    return jsonify(ok=True, days=days)


@app.route('/api/months')
@login_required
def api_months():
    """某类型已有数据的 年-月 列表"""
    dtype = check_type(request.args.get('type', ''))
    root = os.path.join(DATA_DIR, dtype)
    months = []
    if os.path.isdir(root):
        for y in sorted(os.listdir(root)):
            yp = os.path.join(root, y)
            if y.isdigit() and os.path.isdir(yp):
                for m in sorted(os.listdir(yp)):
                    if m.isdigit() and os.path.isdir(os.path.join(yp, m)):
                        n = len([f for f in os.listdir(
                            os.path.join(yp, m)) if f.endswith('.docx')])
                        months.append(
                            {'year': int(y), 'month': int(m), 'count': n})
    return jsonify(ok=True, months=months)


@app.route('/api/record')
@login_required
def api_record():
    dtype = check_type(request.args.get('type', ''))
    y, m, d = parse_date(request.args.get('date', ''))
    path = find_record(dtype, y, m, d)
    if not path:
        return jsonify(ok=False, msg='该日期暂无台账')
    try:
        info = docx_utils.parse_daily(path)
    except Exception as e:
        return jsonify(ok=False, msg=f'解析失败: {e}')
    info['file'] = os.path.basename(path)
    return jsonify(ok=True, record=info)


@app.route('/api/record/image')
@login_required
def api_record_image():
    dtype = check_type(request.args.get('type', ''))
    y, m, d = parse_date(request.args.get('date', ''))
    idx = parse_int(request.args.get('i', 0))
    path = find_record(dtype, y, m, d)
    if not path:
        abort(404)
    res = docx_utils.get_image(path, idx)
    if not res:
        abort(404)
    blob, ctype = res
    return send_file(io.BytesIO(blob), mimetype=ctype)


@app.route('/api/record/save', methods=['POST'])
@login_required
def api_record_save():
    """新增或更新（内容文字）"""
    data = request.get_json(force=True)
    dtype = check_type(data.get('type', ''))
    y, m, d = parse_date(data.get('date', ''))
    # 不允许创建/更新未来日期的台账
    try:
        rec_date = datetime(y, m, d).date()
        if rec_date > datetime.now().date():
            return jsonify(ok=False, msg='不能创建未来日期的台账')
    except ValueError:
        return jsonify(ok=False, msg='日期格式无效')
    content = data.get('content', '')
    sign = data.get('sign', None)
    path = find_record(dtype, y, m, d)
    try:
        if path:
            docx_utils.update_content(path, content, sign=sign)
            msg = '已更新'
        else:
            path = data_path(dtype, y, m, d)
            cfg = config.load_config()
            logo_path = os.path.join(STATIC_IMAGES, cfg['docx_logo']) if cfg.get(
                'docx_logo') else None
            docx_utils.create_daily(path, dtype, y, m, d, content,
                                    company=config.company_full(cfg), logo_path=logo_path)
            msg = '已创建'
    except Exception as e:
        return jsonify(ok=False, msg=f'保存失败: {e}')
    sync_record(dtype, y, m, d, path, content)
    db.log_action(session['user'], '新增台账' if msg == '已创建' else '编辑台账',
                  f'{dtype}/{y:04d}-{m:02d}-{d:02d}',
                  '' if msg == '已创建' else '更新内容', request.remote_addr)
    return jsonify(ok=True, msg=msg, file=os.path.basename(path))


@app.route('/api/record/delete', methods=['POST'])
@login_required
def api_record_delete():
    data = request.get_json(force=True)
    dtype = check_type(data.get('type', ''))
    y, m, d = parse_date(data.get('date', ''))
    path = find_record(dtype, y, m, d)
    if not path:
        return jsonify(ok=False, msg='文件不存在')
    os.remove(path)
    db.delete_record_by_path(rel_path(path))
    db.log_action(session['user'], '删除台账',
                  f'{dtype}/{y:04d}-{m:02d}-{d:02d}', '删除台账文件', request.remote_addr)
    db.sync_filesystem()
    return jsonify(ok=True, msg='已删除')


@app.route('/api/records/delete', methods=['POST'])
@login_required
def api_records_delete():
    """批量删除同一类型下多日台账。body: {type, dates:[YYYY-MM-DD,...]}"""
    data = request.get_json(force=True)
    dtype = check_type(data.get('type', ''))
    dates = data.get('dates') or []
    if not dtype or not dates:
        return jsonify(ok=False, msg='缺少类型或日期')
    ok_cnt = 0
    failed = []
    for ds in dates:
        try:
            y, m, d = parse_date(ds)
            path = find_record(dtype, y, m, d)
            if not path:
                failed.append(ds)
                continue
            os.remove(path)
            db.delete_record_by_path(rel_path(path))
            db.log_action(session['user'], '删除台账',
                          f'{dtype}/{y:04d}-{m:02d}-{d:02d}', '批量删除', request.remote_addr)
            ok_cnt += 1
        except Exception:
            failed.append(ds)
    db.sync_filesystem()
    msg = f'已删除 {ok_cnt} 条'
    if failed:
        msg += f'，{len(failed)} 条失败'
    return jsonify(ok=True, msg=msg, deleted=ok_cnt, failed=failed)


@app.route('/api/records/delete-all', methods=['POST'])
@login_required
def api_records_delete_all():
    """一次性删除（按月维度）。
    body: {type, mode, month?, sm?, em?}
      mode='month' : 删除某月（默认当前月，可传 month='YYYY-MM'）
      mode='range' : 删除月份区间（sm/em 形如 'YYYY-MM'，含两端）
      mode='all'   : 删除该类型全部年份
    """
    data = request.get_json(force=True)
    dtype = check_type(data.get('type', ''))
    if not dtype:
        return jsonify(ok=False, msg='缺少类型')
    mode = data.get('mode', 'all')
    if mode == 'month':
        month = data.get(
            'month') or f"{datetime.now().year}-{datetime.now().month:02d}"
        sy, sm = month.split('-')
        start, end = f'{sy}-{sm}-01', f'{sy}-{sm}-31'
        detail = f'{dtype}/{sy}年{sm}月'
    elif mode == 'range':
        sm, em = data.get('sm'), data.get('em')
        if not sm or not em:
            return jsonify(ok=False, msg='缺少起止月份')
        ssy, ssm = sm.split('-')
        esy, esm = em.split('-')
        start, end = f'{ssy}-{ssm}-01', f'{esy}-{esm}-31'
        detail = f'{dtype}/{sm} 至 {em}'
    elif mode == 'months':
        months = data.get('months') or []
        if not months or not isinstance(months, list):
            return jsonify(ok=False, msg='缺少月份列表')
        all_recs = []
        for mv in months:
            sy, sm = str(mv).split('-')
            recs = db.query_records_range(
                dtype, f'{sy}-{sm}-01', f'{sy}-{sm}-31')
            all_recs.extend(recs)
        detail = f'{dtype}/' + '、'.join(str(m) for m in months)
        return _delete_recs(dtype, all_recs, detail)
    else:
        # 全部年份
        recs_all = db.query_records_range(dtype, '0001-01-01', '9999-12-31')
        return _delete_recs(dtype, recs_all, f'{dtype}（全部年份）')
    recs = db.query_records_range(dtype, start, end)
    return _delete_recs(dtype, recs, detail)


def _delete_recs(dtype, recs, detail):
    if not recs:
        return jsonify(ok=True, msg='没有可删除的台账', deleted=0)
    ok_cnt = 0
    for r in recs:
        try:
            p = os.path.join(BASE, r['path'])
            if os.path.exists(p):
                os.remove(p)
                db.delete_record_by_path(r['path'])
                ok_cnt += 1
        except Exception:
            pass
    db.sync_filesystem()
    db.log_action(session['user'], '删除台账', detail,
                  f'整块删除 {ok_cnt} 条', request.remote_addr)
    return jsonify(ok=True, msg=f'已删除 {ok_cnt} 条（{detail}）', deleted=ok_cnt)


@app.route('/api/record/upload_images', methods=['POST'])
@login_required
def api_upload_images():
    dtype = check_type(request.form.get('type', ''))
    y, m, d = parse_date(request.form.get('date', ''))
    path = find_record(dtype, y, m, d)
    if not path:
        return jsonify(ok=False, msg='请先创建该日台账')
    idx_raw = request.form.get('index', '')
    index = None
    if idx_raw != '':
        try:
            index = int(idx_raw)
        except ValueError:
            index = None
    files = []
    for f in request.files.getlist('images'):
        if f.filename and f.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
            files.append((f.read(), f.filename))
    if not files:
        return jsonify(ok=False, msg='未选择有效图片')
    try:
        used = docx_utils.add_images(path, files, index=index)
    except Exception as e:
        return jsonify(ok=False, msg=f'添加图片失败: {e}')
    if used == 0:
        return jsonify(ok=False, msg='工作图片框已填满（每排4框×2排，共8个），无可放置的空框')
    return jsonify(ok=True, msg=f'已根据照片长宽自动放入 {used} 个照片框')


@app.route('/api/record/download')
@login_required
def api_record_download():
    dtype = check_type(request.args.get('type', ''))
    y, m, d = parse_date(request.args.get('date', ''))
    path = find_record(dtype, y, m, d)
    if not path:
        abort(404)
    return send_file(path, as_attachment=True,
                     download_name=os.path.basename(path))


# ------------------------- 归档 API -------------------------

@app.route('/api/archive/list')
@login_required
def api_archive_list():
    """已归档列表"""
    items = []
    for dtype in [t['name'] for t in db.get_types()]:
        root = os.path.join(BASE, dtype)
        if not os.path.isdir(root):
            continue
        for y in sorted(os.listdir(root)):
            yp = os.path.join(root, y)
            if not (y.isdigit() and os.path.isdir(yp)):
                continue
            months = []
            for m in sorted(os.listdir(yp)):
                mp = os.path.join(yp, m)
                if m.isdigit() and os.path.isdir(mp):
                    n = len([f for f in os.listdir(mp) if f.endswith('.docx')])
                    months.append({'month': int(m), 'count': n})
            items.append({'type': dtype, 'year': int(y), 'months': months})
    return jsonify(ok=True, archives=items)


@app.route('/api/archive/create', methods=['POST'])
@login_required
def api_archive_create():
    """月度归档：把 Data/类型/年/月 复制到 类型/年/月"""
    data = request.get_json(force=True)
    dtype = check_type(data.get('type', ''))
    y = parse_int(data.get('year', 0))
    m = data.get('month')
    if m:
        pairs = [(parse_int(m),)]
    else:
        # 年度归档 = 归档该年所有月份
        yroot = os.path.join(DATA_DIR, dtype, f'{y:04d}')
        if not os.path.isdir(yroot):
            return jsonify(ok=False, msg='该年份无数据')
        pairs = [(int(mm),)
                 for mm in sorted(os.listdir(yroot)) if mm.isdigit()]
    total = 0
    for (mm,) in pairs:
        src = os.path.join(DATA_DIR, dtype, f'{y:04d}', f'{mm:02d}')
        dst = os.path.join(BASE, dtype, f'{y:04d}', f'{mm:02d}')
        if not os.path.isdir(src):
            continue
        os.makedirs(dst, exist_ok=True)
        for f in os.listdir(src):
            if f.endswith('.docx'):
                shutil.copy2(os.path.join(src, f), os.path.join(dst, f))
                total += 1
    if total == 0:
        return jsonify(ok=False, msg='没有可归档的文件')
    return jsonify(ok=True, msg=f'归档完成，共 {total} 个文件')


@app.route('/api/archive/download')
@login_required
def api_archive_download():
    """下载归档压缩包。month 省略则为年度归档"""
    dtype = check_type(request.args.get('type', ''))
    y = parse_int(request.args.get('year', 0))
    m = request.args.get('month')
    if m:
        folder = os.path.join(BASE, dtype, f'{y:04d}', f'{parse_int(m):02d}')
        name = f'{dtype}-{y:04d}-{int(m):02d}.zip'
        arc = f'{dtype}/{y:04d}/{int(m):02d}'
    else:
        folder = os.path.join(BASE, dtype, f'{y:04d}')
        name = f'{dtype}-{y:04d}.zip'
        arc = f'{dtype}/{y:04d}'
    if not os.path.isdir(folder):
        abort(404)
    buf = zip_folder_to_memory(folder, arc)
    return send_file(buf, as_attachment=True, download_name=name,
                     mimetype='application/zip')


@app.route('/api/archive/delete', methods=['POST'])
@login_required
def api_archive_delete():
    """删除归档：按 类型/年/月 删除单个归档目录，或留空月份删除整年归档。
    仅删除归档目录(BASE/...)，不影响 Data 源数据。"""
    data = request.get_json(force=True) if request.is_json else request.form
    dtype = check_type(data.get('type', ''))
    y = parse_int(data.get('year', 0))
    m = data.get('month')
    if m:
        folder = os.path.join(BASE, dtype, f'{y:04d}', f'{parse_int(m):02d}')
        label = f'{dtype}/{y:04d}/{int(m):02d}'
    else:
        folder = os.path.join(BASE, dtype, f'{y:04d}')
        label = f'{dtype}/{y:04d}'
    if not os.path.isdir(folder):
        return jsonify(ok=False, msg='归档目录不存在，可能已被删除')
    try:
        shutil.rmtree(folder)
    except Exception as e:
        return jsonify(ok=False, msg=f'删除失败: {e}')
    return jsonify(ok=True, msg=f'已删除归档 {label}')


# ------------------------- 类型管理 / 搜索 -------------------------

@app.route('/api/types', methods=['GET'])
@login_required
def api_types():
    return jsonify(ok=True, types=db.get_types())


@app.route('/api/types', methods=['POST'])
@login_required
def api_type_add():
    name = (request.get_json(force=True) or {}).get('name', '')
    ok, res = db.add_type(name)
    if not ok:
        return jsonify(ok=False, msg=res)
    db.log_action(session['user'], '新增类型', name, '', request.remote_addr)
    return jsonify(ok=True, id=res, name=name)


@app.route('/api/types/<int:tid>', methods=['PUT'])
@login_required
def api_type_rename(tid):
    name = (request.get_json(force=True) or {}).get('name', '')
    ok, res = db.rename_type(tid, name)
    if not ok:
        return jsonify(ok=False, msg=res)
    db.log_action(session['user'], '重命名类型',
                  f'{tid}→{name}', res, request.remote_addr)
    return jsonify(ok=True, id=tid, name=name)


@app.route('/api/types/<int:tid>', methods=['DELETE'])
@login_required
def api_type_delete(tid):
    ok, res = db.delete_type(tid)
    if not ok:
        return jsonify(ok=False, msg=res)
    db.log_action(session['user'], '删除类型', res, '', request.remote_addr)
    return jsonify(ok=True, msg=f'已删除类型「{res}」')


# 兼容前端旧路径（/api/type/* 别名指向 /api/types/*）
@app.route('/api/type/add', methods=['POST'])
@login_required
def api_type_add_alias():
    return api_type_add()


@app.route('/api/type/rename/<int:tid>', methods=['POST'])
@login_required
def api_type_rename_alias(tid):
    return api_type_rename(tid)


@app.route('/api/type/<int:tid>', methods=['DELETE'])
@login_required
def api_type_delete_alias(tid):
    return api_type_delete(tid)


@app.route('/api/search')
@login_required
def api_search():
    kw = request.args.get('kw', '').strip()
    if not kw:
        return jsonify(ok=True, results=[])
    return jsonify(ok=True, results=db.search(kw))


@app.route('/api/stats')
@login_required
def api_stats():
    """数据大屏统计：汇总卡片、各类型数量、近12月趋势、最近记录。"""
    return jsonify(
        ok=True,
        totals=db.stats_totals(),
        by_type=db.stats_by_type(),
        monthly=db.stats_monthly(),
        recent=db.recent_records(10),
    )


# ------------------------- 智能总结 -------------------------

def _split_items(content):
    """把一段台账内容拆成独立工作条目（按行，去掉序号/项目符号）。"""
    items = []
    for line in (content or '').splitlines():
        line = line.strip()
        if not line:
            continue
        line = re.sub(r'^[\d]+[、.．)）\]\s]+', '', line)
        line = re.sub(r'^[•·\-–\*\u2022\s]+', '', line)
        if line:
            items.append(line)
    return items


def _keywords(text, top=10):
    """中文双字粒度的关键词（无分词依赖），仅取出现>=2次的。"""
    text = re.sub(r'[\s\W_0-9a-zA-Z]+', '', text or '')
    cnt = collections.Counter(text[i:i + 2] for i in range(len(text) - 1))
    return [w for w, c in cnt.most_common(top) if c >= 2]


_PERIOD_LABELS = {
    'year': '全年', 'half1': '上半年（1-6月）', 'half2': '下半年（7-12月）',
    'q1': '第一季度', 'q2': '第二季度', 'q3': '第三季度', 'q4': '第四季度',
}


def _period_label(period, year):
    if period in _PERIOD_LABELS:
        return _PERIOD_LABELS[period]
    if period.startswith('m'):
        return f'{int(period[1:])}月'
    return period


def _period_range(period, year):
    """返回 (start, end) 日期字符串。"""
    if period == 'year':
        return f'{year}-01-01', f'{year}-12-31'
    if period == 'half1':
        return f'{year}-01-01', f'{year}-06-30'
    if period == 'half2':
        return f'{year}-07-01', f'{year}-12-31'
    if period.startswith('q'):
        q = int(period[1])
        sm, em = (q - 1) * 3 + 1, q * 3
        return (f'{year}-{sm:02d}-01',
                f'{year}-{em:02d}-{calendar.monthrange(year, em)[1]}')
    if period.startswith('m'):
        m = int(period[1:])
        return (f'{year}-{m:02d}-01',
                f'{year}-{m:02d}-{calendar.monthrange(year, m)[1]}')
    return None, None


@app.route('/api/summary')
@login_required
def api_summary():
    """按周期自动总结台账：月度 / 季度 / 半年 / 全年 / 自定义范围。
    参数：type（类型名或'全部'）、year、period（year/half1/half2/q1..q4/m1..m12/range），
    自定义范围时额外传 start、end（YYYY-MM-DD）。"""
    type_name = request.args.get('type', '全部')
    period = request.args.get('period', 'year')
    if period == 'range':
        start = request.args.get('start', '')
        end = request.args.get('end', '')
        if not start or not end:
            return jsonify(ok=False, msg='请选择开始和结束日期'), 400
        if start > end:
            start, end = end, start
        label = f'{start} 至 {end}'
        try:
            year = int(start[:4])
        except ValueError:
            year = datetime.now().year
    else:
        year = parse_int(request.args.get('year', datetime.now().year),
                         default=datetime.now().year)
        start, end = _period_range(period, year)
        if not start:
            return jsonify(ok=False, msg='未知的周期类型'), 400
        label = _period_label(period, year)

    return build_summary(type_name, year, period, start, end, label)


def build_summary(type_name, year, period, start, end, label):
    """根据周期聚合台账，生成智能总结数据（供 /api/summary 与导出共用）。"""
    rows = db.query_records_range(
        None if type_name == '全部' else type_name, start, end)
    if not rows:
        empty_year = '' if period == 'range' else f'{year}年'
        return dict(ok=True, empty=True, label=label,
                    type=type_name, year=year, period=period,
                    narrative=f'{empty_year}{label}暂无台账数据。')

    record_count = len(rows)
    dates = sorted({r['date'] for r in rows})
    days = len(dates)
    all_items = []
    for r in rows:
        all_items.extend(_split_items(r['content']))
    item_count = len(all_items)
    by_type = collections.Counter(r['type_name'] for r in rows)
    by_month = collections.Counter(r['date'][:7] for r in rows)
    kw = _keywords('\n'.join(r['content'] for r in rows))
    reps = []
    for r in sorted(rows, key=lambda x: x['date']):
        its = _split_items(r['content'])
        if its:
            txt = re.sub(r'[；;。，,、\s]+$', '', its[0])
            reps.append({'date': r['date'], 'text': txt})
        if len(reps) >= 6:
            break

    scope = f'「{type_name}」' if type_name != '全部' else ''
    head = f'在{label}期间，{scope}' if period == 'range' else f'在{year}年{label}期间，{scope}'
    parts = [head + f'共录入台账 {record_count} 份，'
             f'覆盖 {days} 天，累计记录工作内容 {item_count} 项。']
    if type_name == '全部' and by_type:
        parts.append(
            '各类型分布：' + '、'.join(f'{k}{v}份' for k, v in by_type.most_common()) + '。')
    if kw:
        parts.append('高频关键词：' + '、'.join(kw) + '。')
    if reps:
        parts.append(
            '重点工作示例：' + '；'.join(f"{x['date']} {x['text']}" for x in reps[:4]) + '。')
    narrative = ''.join(parts)

    return dict(
        ok=True, empty=False, label=label, type=type_name, year=year, period=period,
        record_count=record_count, days=days, item_count=item_count,
        by_type=dict(by_type), by_month=dict(by_month),
        keywords=kw, representatives=reps, narrative=narrative,
    )


@app.route('/api/summary/export')
@login_required
def api_summary_export():
    """导出智能总结为 Word 文档。"""
    type_name = request.args.get('type', '全部')
    period = request.args.get('period', 'year')
    if period == 'range':
        start = request.args.get('start', '')
        end = request.args.get('end', '')
        if not start or not end:
            return jsonify(ok=False, msg='请选择开始和结束日期'), 400
        if start > end:
            start, end = end, start
        label = f'{start} 至 {end}'
        try:
            year = int(start[:4])
        except ValueError:
            year = datetime.now().year
    else:
        year = parse_int(request.args.get('year', datetime.now().year),
                         default=datetime.now().year)
        start, end = _period_range(period, year)
        if not start:
            return jsonify(ok=False, msg='未知的周期类型'), 400
        label = _period_label(period, year)

    s = build_summary(type_name, year, period, start, end, label)
    if s.get('empty'):
        return jsonify(ok=False, msg='该周期暂无数据，无法导出'), 400

    doc = Document()
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal'].font.size = Pt(11)
    # 标题
    doc.add_heading(s['label'], level=0)
    sub = doc.add_paragraph()
    run = sub.add_run(('「%s」' % s['type']) if s['type'] != '全部' else '全部类型')
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    # 统计概览
    doc.add_heading('数据概览', level=1)
    stats = [
        ('录入台账', '%d 份' % s['record_count']),
        ('覆盖天数', '%d 天' % s['days']),
        ('记录事项', '%d 项' % s['item_count']),
    ]
    if s.get('keywords'):
        stats.append(('高频关键词', '、'.join(s['keywords'])))
    if s.get('by_type') and s['type'] == '全部':
        stats.append(
            ('各类型分布', '、'.join(f'{k}{v}份' for k, v in s['by_type'].items())))
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Light Grid Accent 1'
    for k, v in stats:
        row = tbl.add_row().cells
        row[0].text = k
        row[1].text = v
    # 正文
    doc.add_heading('智能总结', level=1)
    for para in s['narrative'].split('\n'):
        if para.strip():
            doc.add_paragraph(para.strip())
    # 重点工作示例
    if s.get('representatives'):
        doc.add_heading('重点工作示例', level=1)
        for x in s['representatives'][:6]:
            doc.add_paragraph(f"{x['date']}：{x['text']}", style='List Bullet')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    raw = '智能总结_%s_%s.docx' % (s['type'] if s['type'] != '全部' else '全部',
                               s['label'].replace(' ', '').replace('/', '-'))
    from urllib.parse import quote
    # ASCII 备用名 + RFC5987 中文名（filename* 优先级更高，真实浏览器显示中文）
    cd = "attachment; filename=\"summary.docx\"; filename*=UTF-8''%s" % quote(
        raw)
    resp = send_file(
        buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    resp.headers['Content-Disposition'] = cd
    return resp


# ------------------------- 重建索引 -------------------------

@app.route('/api/index/rebuild', methods=['POST'])
@login_required
def api_index_rebuild():
    """重建台账索引：重新扫描文件系统并重建 SQLite 索引。"""
    try:
        cnt = db.rebuild_index()
    except Exception as e:
        return jsonify(ok=False, msg=f'重建索引失败: {e}')
    db.log_action(session['user'], '重建索引', '系统',
                  f'重建后共 {cnt} 条记录', request.remote_addr)
    return jsonify(ok=True, msg=f'索引已重建，共 {cnt} 条记录', count=cnt)


# ------------------------- 免费 AI 总结（本地规则，无需 KEY） -------------------------

@app.route('/api/summary/free')
@login_required
def api_summary_free():
    """免费智能总结：复用本地规则生成，无需任何 API KEY。"""
    type_name = request.args.get('type', '全部')
    period = request.args.get('period', 'year')
    res = build_summary(type_name, _summary_year(), period,
                        *_summary_range(period), _summary_label(period))
    if not res.get('empty'):
        db.log_action(session['user'], '智能总结-免费',
                      type_name, period, request.remote_addr)
    return jsonify(res)


# ------------------------- 远程 AI 总结（需 KEY） -------------------------

@app.route('/api/summary/ai', methods=['POST'])
@login_required
def api_summary_ai():
    """调用用户配置的远程大模型生成总结。需先在系统设置填写 AI Key。"""
    cfg = config.load_config()
    ai = cfg.get('ai', {})
    api_key = ai.get('api_key', '').strip()
    base_url = ai.get('base_url', '').strip()
    model = ai.get('model', '').strip()
    if not api_key or not base_url or not model:
        return jsonify(ok=False, msg='尚未配置 AI：请前往「系统设置 → AI 配置」填写 API 地址、Key 与模型')
    data = request.get_json(force=True, silent=True) or {}
    type_name = data.get('type', '全部')
    period = data.get('period', 'year')
    s = build_summary(type_name, _summary_year(), period,
                      *_summary_range(period), _summary_label(period))
    if s.get('empty'):
        return jsonify(ok=True, empty=True, narrative=s['narrative'])
    prompt = (
        '你是一名物业工作台账分析助手。请基于以下台账数据，用中文生成一段专业、'
        '条理清晰的总结，包含工作概况、重点事项、趋势与建议，不超过 400 字。\n\n'
        f'范围：{s["label"]}（{type_name}）\n'
        f'录入 {s["record_count"]} 份，覆盖 {s["days"]} 天，记录 {s["item_count"]} 项。\n'
        + (f'各类型分布：{", ".join(f"{k}{v}份" for k,
           v in s["by_type"].items())}\n' if s.get('by_type') and type_name == '全部' else '')
        + (f'高频关键词：{", ".join(s["keywords"])}\n' if s.get('keywords') else '')
        + '重点工作示例：' +
        '；'.join(f'{x["date"]} {x["text"]}' for x in s.get(
            'representatives', [])[:5]) + '\n'
    )
    try:
        text = _call_ai(base_url, api_key, model, prompt)
    except Exception as e:
        return jsonify(ok=False, msg=f'AI 调用失败: {e}')
    db.log_action(session['user'], '智能总结-AI', type_name,
                  period, request.remote_addr)
    return jsonify(ok=True, empty=False, narrative=text, base=s)


def _call_ai(base_url, api_key, model, prompt):
    """调用 OpenAI 兼容的 chat/completions 接口，返回生成的文本。"""
    import urllib.request
    url = base_url.rstrip('/') + '/chat/completions'
    payload = {
        'model': model,
        'messages': [{'role': 'user', 'content': prompt}],
        'temperature': 0.7,
        'max_tokens': 800,
    }
    req = urllib.request.Request(
        url, data=json.dumps(payload).encode('utf-8'),
        headers={'Authorization': f'Bearer {api_key}',
                 'Content-Type': 'application/json'})
    with urllib.request.urlopen(req, timeout=60) as resp:
        data = json.loads(resp.read().decode('utf-8'))
    return data['choices'][0]['message']['content'].strip()


@app.route('/api/ai/models')
@login_required
def api_ai_models():
    """拉取用户配置的 AI 平台可用模型列表（OpenAI 兼容 /models 接口）。
    返回每个模型的 id 及中文简介（针对已知免费模型做标注）。"""
    cfg = config.load_config()
    ai = cfg.get('ai', {})
    base_url = (ai.get('base_url') or '').strip()
    api_key = (ai.get('api_key') or '').strip()
    if not base_url or not api_key:
        return jsonify(ok=False, msg='请先在「系统设置 → AI 配置」填写 API 地址和 Key，再拉取模型')
    try:
        import urllib.request
        url = base_url.rstrip('/') + '/models'
        req = urllib.request.Request(
            url, headers={'Authorization': f'Bearer {api_key}',
                          'Accept': 'application/json'})
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read().decode('utf-8'))
    except Exception as e:
        return jsonify(ok=False, msg=f'拉取模型失败: {e}')
    raw = data.get('data', [])
    models = [{'id': m.get('id'), 'desc': _model_desc(
        m.get('id', ''))} for m in raw]
    return jsonify(ok=True, models=models, provider=ai.get('provider', ''))


_MODEL_DESC = {
    'Qwen/Qwen3-8B-Instruct': '阿里通义千问 Qwen3-8B：开源免费模型，中文能力优秀，适合日常台账总结，速度快、成本低。',
    'Qwen/Qwen2.5-7B-Instruct': '阿里通义千问 2.5-7B：稳定可靠的中文对话模型，适合通用总结。',
    'deepseek-ai/DeepSeek-V3': '深度求索 DeepSeek-V3：强推理大模型，适合复杂分析与长文本总结。',
    'deepseek-ai/DeepSeek-R1': '深度求索 DeepSeek-R1：推理增强模型，擅长分步思考与深度分析。',
    'THUDM/glm-4-9b-chat': '智谱 GLM-4-9B：中文表现均衡，适合总结与问答。',
    'meta-llama/Llama-3.3-70B-Instruct': 'Meta Llama-3.3-70B：英文强、中文良好，适合多语言场景。',
}


def _model_desc(mid):
    if mid in _MODEL_DESC:
        return _MODEL_DESC[mid]
    if 'qwen' in mid.lower():
        return '通义千问系列模型，中文能力强，适合台账总结。'
    if 'deepseek' in mid.lower():
        return 'DeepSeek 系列模型，推理能力突出。'
    if 'glm' in mid.lower():
        return '智谱 GLM 系列，中文表现均衡。'
    return '该模型可用于生成总结，具体能力请参考平台说明。'


def _summary_year():
    return parse_int(request.args.get('year', datetime.now().year),
                     default=datetime.now().year)


def _summary_label(period):
    if period == 'range':
        return f'{request.args.get("start", "")} 至 {request.args.get("end", "")}'
    return _PERIOD_LABELS.get(period, period)


def _summary_range(period):
    if period == 'range':
        s = request.args.get('start', '')
        e = request.args.get('end', '')
        return s, e
    return _period_range(period, _summary_year())


# ------------------------- 系统日志 -------------------------

@app.route('/logs')
@login_required
def logs_page():
    return _render('logs.html')


@app.route('/api/logs')
@login_required
def api_logs():
    """查询系统操作日志，支持 action / user / start / end 筛选。"""
    action = request.args.get('action', '').strip() or None
    user = request.args.get('user', '').strip() or None
    start = request.args.get('start', '').strip() or None
    end = request.args.get('end', '').strip() or None
    limit = parse_int(request.args.get('limit', 200), default=200)
    rows = db.recent_logs(limit=limit, action=action,
                          user=user, start=start, end=end)
    return jsonify(ok=True, logs=rows, total=db.count_logs())


@app.route('/api/logs/delete', methods=['POST'])
@login_required
def api_logs_delete():
    """删除系统日志：支持按 action/user/start/end 条件删除，或 action 传 'all' 清空整表。"""
    data = request.get_json(silent=True) or {}
    mode = data.get('mode', 'range')
    action = (data.get('action') or '').strip() or None
    user = (data.get('user') or '').strip() or None
    start = (data.get('start') or '').strip() or None
    end = (data.get('end') or '').strip() or None
    if mode == 'all':
        n = db.delete_logs()
        detail = '全部系统日志'
    else:
        if not (action or user or start or end):
            return jsonify(ok=False, msg='请设置删除条件（操作类型 / 用户 / 起止日期）')
        if end:
            end = end + ' 23:59:59'
        n = db.delete_logs(action=action, user=user, start=start, end=end)
        parts = []
        if action:
            parts.append('操作=' + action)
        if user:
            parts.append('用户=' + user)
        if start:
            parts.append('起=' + start)
        if end:
            parts.append('止=' + end.replace(' 23:59:59', ''))
        detail = '、'.join(parts) + ' 的系统日志'
    db.log_action(session.get('user'), '删除日志', detail,
                  f'删除 {n} 条', request.remote_addr)
    return jsonify(ok=True, deleted=n, total=db.count_logs())


# ------------------------- 导入 / 导出 -------------------------

def _import_one_file(f, default_type):
    """导入单个文件，返回结果字典（含 dup 标记是否覆盖已有记录）。"""
    fname = f.filename or ''
    if not fname.lower().endswith('.docx'):
        return {'file': fname, 'ok': False, 'msg': '仅支持 .docx', 'dup': False}
    mname = re.match(r'^\d+\.(.+?)\d{4}-\d{2}-\d{2}\.docx$', fname)
    dtype = mname.group(1) if (mname and db.type_id(
        mname.group(1))) else default_type
    tmp = os.path.join(BASE, '_import_tmp.docx')
    f.save(tmp)
    try:
        mt = FNAME_RE.match(os.path.basename(fname))
        n_tables = docx_utils.count_tables(tmp)
        if n_tables <= 1:
            # 单日文件
            info = docx_utils.parse_daily(tmp)
            date = None
            if mt:
                date = (int(mt.group(3)), int(mt.group(4)), int(mt.group(5)))
            elif info and info['date']:
                y_, m_, d_ = info['date'].split('-')
                date = (int(y_), int(m_), int(d_))
            if not date:
                return {'file': fname, 'ok': False, 'msg': '无法识别日期', 'dup': False}
            y, m, d = date
            dst = data_path(dtype, y, m, d)
            existed = find_record(dtype, y, m, d) is not None
            os.makedirs(os.path.dirname(dst), exist_ok=True)
            shutil.copy2(tmp, dst)
            content = (info or {}).get('content', '') if info else ''
            sync_record(dtype, y, m, d, dst, content)
            return {'file': fname, 'ok': True, 'msg': f'导入为 {os.path.basename(dst)}',
                    'dup': existed}
        else:
            # 月度文件 -> 拆分
            doc_dates = docx_utils.split_monthly_docx(
                tmp, dtype,
                [os.path.join(DATA_DIR, dtype, '_tmp_split')])
            tmp_dir = os.path.join(DATA_DIR, dtype, '_tmp_split')
            n = 0
            dup_n = 0
            for (y, m, d), fn in doc_dates:
                src = os.path.join(tmp_dir, fn)
                dst = data_path(dtype, y, m, d)
                existed = find_record(dtype, y, m, d) is not None
                os.makedirs(os.path.dirname(dst), exist_ok=True)
                shutil.move(src, dst)
                sync_record(dtype, y, m, d, dst, '')
                n += 1
                if existed:
                    dup_n += 1
            shutil.rmtree(tmp_dir, ignore_errors=True)
            return {'file': fname, 'ok': True,
                    'msg': f'月度文件已拆分导入 {n} 天' + (f'，其中 {dup_n} 天为覆盖' if dup_n else ''),
                    'dup': dup_n > 0}
    except Exception as e:
        return {'file': fname, 'ok': False, 'msg': f'导入失败: {e}', 'dup': False}
    finally:
        if os.path.exists(tmp):
            os.remove(tmp)


@app.route('/api/import', methods=['POST'])
@login_required
def api_import():
    """批量导入：每日 docx（按文件名识别）或月度 docx（自动拆分）。
    支持多文件（files 字段）或单文件；前端逐文件请求时每个文件单独返回。
    需指定类型；若文件名含 安保/消控 则以文件名为准。
    返回结果含 dup 标记（True 表示该记录覆盖了已存在的台账）。"""
    default_type = request.form.get('type', '').strip()
    if not default_type or db.type_id(default_type) is None:
        return jsonify(ok=False, msg='请选择有效的目标类型')
    files = request.files.getlist('files')
    # 兼容单文件字段名 file
    if not files:
        files = request.files.getlist('file')
    if not files:
        return jsonify(ok=False, msg='未收到文件')
    results = [_import_one_file(f, default_type) for f in files]
    db.sync_filesystem()
    db.log_action(session['user'], '导入', default_type,
                  f'导入 {len(results)} 个文件', request.remote_addr)
    return jsonify(ok=True, results=results)


@app.route('/api/export')
@login_required
def api_export():
    """批量导出：type=安保|消控|全部, year 必填, month 可选, kind=zip|year"""
    dtype = request.args.get('type', '全部')
    y = parse_int(request.args.get('year', 0))
    m = request.args.get('month')
    kind = request.args.get('kind', 'zip')
    if dtype == '全部':
        types = [t['name'] for t in db.get_types()]
    else:
        check_type(dtype)
        types = [dtype]
    buf = io.BytesIO()
    total = 0
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for t in types:
            if m:
                folders = [(os.path.join(DATA_DIR, t, f'{y:04d}', f'{int(m):02d}'),
                            f'{t}/{y:04d}/{int(m):02d}')]
            else:
                yroot = os.path.join(DATA_DIR, t, f'{y:04d}')
                folders = []
                if os.path.isdir(yroot):
                    for mm in sorted(os.listdir(yroot)):
                        if mm.isdigit():
                            folders.append(
                                (os.path.join(yroot, mm), f'{t}/{y:04d}/{mm}'))
            for folder, arc in folders:
                if not os.path.isdir(folder):
                    continue
                for f in sorted(os.listdir(folder)):
                    if f.endswith('.docx'):
                        zf.write(os.path.join(folder, f), f'{arc}/{f}')
                        total += 1
    if total == 0:
        return jsonify(ok=False, msg='所选范围没有数据'), 404
    buf.seek(0)
    db.log_action(session['user'], '导出', f'{dtype}/{y:04d}' + (f'-{int(m):02d}' if m else ''),
                  f'导出 {total} 个文件', request.remote_addr)
    name = f'台账导出-{dtype}-{y:04d}' + (f'-{int(m):02d}' if m else '') + '.zip'
    if kind == 'year' and not m:
        name = f'年度总台账-{dtype}-{y:04d}.zip'
    return send_file(buf, as_attachment=True, download_name=name,
                     mimetype='application/zip')


# ------------------------- 系统备份 / 恢复 -------------------------

BACKUP_DIR = os.path.join(BASE, 'Backup')


def _backup_members():
    """返回备份内容清单 [(磁盘绝对路径, 压缩包内相对路径)]。
    备份范围：DB/（数据库）、users.json（账号）、Data/（台账文档）。"""
    members = []
    if os.path.isdir(db.DB_DIR):
        for f in sorted(os.listdir(db.DB_DIR)):
            full = os.path.join(db.DB_DIR, f)
            if os.path.isfile(full):
                members.append((full, f'DB/{f}'))
    if os.path.isfile(USERS_FILE):
        members.append((USERS_FILE, 'users.json'))
    for root, _dirs, files in os.walk(DATA_DIR):
        for f in files:
            full = os.path.join(root, f)
            rel = os.path.relpath(full, BASE).replace('\\', '/')
            members.append((full, rel))
    return members


@app.route('/api/backup/download')
@login_required
def api_backup_download():
    """下载完整系统备份（DB/ + users.json + Data/）"""
    members = _backup_members()
    if not members:
        return jsonify(ok=False, msg='暂无可备份的数据'), 404
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for full, arc in members:
            zf.write(full, arc)
    buf.seek(0)
    name = f'系统备份-{datetime.now().strftime("%Y%m%d-%H%M%S")}.zip'
    return send_file(buf, as_attachment=True, download_name=name,
                     mimetype='application/zip')


@app.route('/api/backup/info')
@login_required
def api_backup_info():
    """备份范围概览"""
    members = _backup_members()
    size = sum(os.path.getsize(f) for f, _ in members if os.path.exists(f))
    return jsonify(ok=True,
                   db_file=os.path.relpath(
                       db.DB_PATH, BASE).replace('\\', '/'),
                   files=len(members),
                   size_mb=round(size / 1048576, 2),
                   scope=['DB/（SQLite 数据库）', 'users.json（账号）', 'Data/（台账文档）'])


@app.route('/api/backup/restore', methods=['POST'])
@login_required
def api_backup_restore():
    """从备份压缩包恢复；恢复前自动把现有数据另存到 Backup/ 目录"""
    up = request.files.get('file')
    if not up or not up.filename.lower().endswith('.zip'):
        return jsonify(ok=False, msg='请选择备份 .zip 文件')
    os.makedirs(BACKUP_DIR, exist_ok=True)
    safety = os.path.join(
        BACKUP_DIR, f'restore-before-{datetime.now().strftime("%Y%m%d-%H%M%S")}.zip')
    with zipfile.ZipFile(safety, 'w', zipfile.ZIP_DEFLATED) as zf:
        for full, arc in _backup_members():
            zf.write(full, arc)

    buf = io.BytesIO(up.read())
    restored = 0
    base_abs = os.path.abspath(BASE)
    try:
        with zipfile.ZipFile(buf) as zf:
            for info in zf.infolist():
                if info.is_dir():
                    continue
                # 统一为 '/' 分隔并去除前导斜杠
                arc = info.filename.replace('\\', '/').lstrip('/')
                # 仅允许恢复 DB / Data / users.json 三类成员
                top = arc.split('/')[0]
                if top not in ('DB', 'Data') and arc != 'users.json':
                    continue
                # 防 zip slip：规范化后必须仍位于 BASE 内（阻止 ..\ 等绕过）
                dst = os.path.abspath(os.path.join(BASE, arc))
                if os.path.commonpath([base_abs, dst]) != base_abs:
                    continue
                os.makedirs(os.path.dirname(dst), exist_ok=True)
                with zf.open(info) as src, open(dst, 'wb') as out:
                    shutil.copyfileobj(src, out)
                restored += 1
    except zipfile.BadZipFile:
        return jsonify(ok=False, msg='备份文件已损坏或不是有效的 zip')
    db.init_db()
    load_users()
    db.log_action(session['user'], '备份恢复', '系统',
                  f'恢复 {restored} 个文件', request.remote_addr)
    return jsonify(ok=True, msg=f'恢复完成，共写入 {restored} 个文件；'
                   f'原数据已另存为 Backup/{os.path.basename(safety)}')


# ------------------------- 系统设置 API -------------------------

@app.route('/api/config', methods=['GET', 'POST'])
@login_required
def api_config():
    if request.method == 'GET':
        ver, _ = load_version()
        data = config.load_config()
        data['version'] = ver
        return jsonify(ok=True, config=data)
    # 保存系统设置（支持 multipart 表单，含 LOGO 上传）
    form = request.form
    data = config.load_config()
    data['system_title'] = (form.get('system_title')
                            or '').strip() or data['system_title']
    data['company_short_name'] = (form.get('company_short_name')
                                  or '').strip() or data['company_short_name']
    data['customer_name'] = (form.get('customer_name') or '').strip()
    # LOGO 上传
    logo = request.files.get('logo') if 'logo' in request.files else None
    if logo and logo.filename:
        ext = os.path.splitext(logo.filename)[1].lower()
        if ext not in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.svg', '.webp'):
            return jsonify(ok=False, msg='不支持的 LOGO 格式')
        os.makedirs(STATIC_IMAGES, exist_ok=True)
        fname = 'company_logo' + ext
        logo.save(os.path.join(STATIC_IMAGES, fname))
        data['company_logo'] = '/static/Images/' + fname
        data['docx_logo'] = fname
    saved = config.save_config(data)
    db.log_action(session['user'], '系统设置', '系统', '保存系统设置', request.remote_addr)
    return jsonify(ok=True, msg='设置已保存', config=saved)


@app.route('/api/ai/config', methods=['GET', 'POST'])
@login_required
def api_ai_config():
    """AI 配置读写：provider / base_url / api_key / model。"""
    if request.method == 'GET':
        cfg = config.load_config()
        return jsonify(ok=True, ai=cfg.get('ai', {}))
    data = request.get_json(force=True, silent=True) or {}
    cur = config.load_config()
    ai = dict(cur.get('ai', {}))
    for k in ('provider', 'base_url', 'api_key', 'model'):
        if k in data:
            ai[k] = data[k]
    cur['ai'] = ai
    saved = config.save_config(cur)
    db.log_action(session['user'], '系统设置', 'AI配置',
                  '保存 AI 配置', request.remote_addr)
    return jsonify(ok=True, msg='AI 配置已保存', ai=saved['ai'])


@app.route('/api/version')
@login_required
def api_version():
    ver, history = load_version()
    return jsonify(ok=True, version=ver, history=history)


@app.route('/api/about')
@login_required
def api_about():
    cfg = config.load_config()
    counts = {r['name']: r['count'] for r in db.stats_by_type()}
    ver, _ = load_version()
    return jsonify(ok=True, about={
        'version': ver,
        'system_title': cfg['system_title'],
        'company_short_name': cfg['company_short_name'],
        'customer_name': cfg['customer_name'],
        'company_logo': cfg['company_logo'],
        'company_full': config.company_full(cfg),
        'record_counts': counts,
        'current_user': session.get('user'),
    })


@app.route('/api/current_user')
def api_current_user():
    """返回当前 Web 端登录的用户名（桌面端轮询用）。"""
    return jsonify(ok=True, user=session.get('user'))


@app.route('/api/account/password', methods=['POST'])
@login_required
def api_account_password():
    data = request.get_json(force=True, silent=True) or {}
    ok, msg = change_password(
        session['user'], data.get('old', ''), data.get('new', ''))
    if ok:
        db.log_action(session['user'], '修改密码',
                      session['user'], '', request.remote_addr)
    return jsonify(ok=ok, msg=msg)


@app.route('/api/users', methods=['GET', 'POST'])
@login_required
def api_users():
    if request.method == 'GET':
        return jsonify(ok=True, users=list_user_names())
    data = request.get_json(force=True, silent=True) or {}
    ok, msg = add_user((data.get('username') or '').strip(),
                       data.get('password', ''))
    if ok:
        db.log_action(session['user'], '新增用户', data.get(
            'username', ''), '', request.remote_addr)
    return jsonify(ok=ok, msg=msg)


@app.route('/api/users/<name>', methods=['DELETE'])
@login_required
def api_user_delete(name):
    ok, msg = delete_user(name)
    if ok:
        db.log_action(session['user'], '删除用户', name, '', request.remote_addr)
    return jsonify(ok=ok, msg=msg)


@app.route('/api/users/<name>/reset', methods=['POST'])
@login_required
def api_user_reset(name):
    data = request.get_json(force=True, silent=True) or {}
    ok, msg = reset_password(name, data.get('password', ''))
    if ok:
        db.log_action(session['user'], '重置密码', name, '', request.remote_addr)
    return jsonify(ok=ok, msg=msg)


def run_server():
    """初始化并启动 Flask 服务（供 __main__ 与 PyInstaller worker 子进程共用）。"""
    import time as _t
    _t0 = _t.time()
    os.makedirs(DATA_DIR, exist_ok=True)
    os.makedirs(db.DB_DIR, exist_ok=True)
    load_users()
    print('[启动] 用户加载完成 (%.2fs)' % (_t.time() - _t0))
    db.init_db()
    print('[启动] 数据库初始化完成 (%.2fs)' % (_t.time() - _t0))
    # 关闭 debug/reloader：避免 fork 子进程导致 init_db 与全量文件扫描执行两遍，
    # 这是启动慢的主因。如需热重载开发，可临时设 debug=True。
    app.run(host='0.0.0.0', port=8088, debug=False, use_reloader=False)


if __name__ == '__main__':
    run_server()
