# -*- coding: utf-8 -*-
"""将月度台账拆分为每日文件。
输出到 [类型]/[年份]/[月份]/ 以及 Data/[类型]/[年份]/[月份]/（供系统使用）。
用法: python split_monthly.py 安保7月份.docx 安保
"""
import os
import sys
from docx_utils import split_monthly_docx, find_date_in_table
from docx import Document

BASE = os.path.dirname(os.path.abspath(__file__))


def split(src, dtype):
    doc = Document(src)
    if not doc.tables:
        print(f'[跳过] {src} 无表格')
        return
    date = find_date_in_table(doc.tables[0])
    if not date:
        print(f'[跳过] {src} 未在表格中找到日期')
        return
    y, m = date[0], date[1]
    out1 = os.path.join(BASE, dtype, f'{y:04d}', f'{m:02d}')
    out2 = os.path.join(BASE, 'Data', dtype, f'{y:04d}', f'{m:02d}')
    results = split_monthly_docx(src, dtype, [out1, out2])
    print(f'{src} -> 共拆分 {len(results)} 天')
    for (yy, mm, dd), fname in results:
        print(f'  {fname}')


if __name__ == '__main__':
    if len(sys.argv) >= 3:
        split(sys.argv[1], sys.argv[2])
    else:
        split(os.path.join(BASE, '安保7月份.docx'), '安保')
        split(os.path.join(BASE, '消控7月份.docx'), '消控')
