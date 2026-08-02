# -*- coding: utf-8 -*-
"""乾明工作台账系统 - Word 文档处理工具
功能：月度台账拆分为每日文件、每日台账解析/编辑/新建、图片提取。
"""
import io
import os
import re
import shutil
from PIL import Image as PILImage
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.opc.constants import RELATIONSHIP_TYPE as RT

W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
DATE_RE = re.compile(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日')

TYPES = ('安保', '消控')
DEPT_NAME = {'安保': '安保部', '消控': '安保部（消控）'}

# 公司全称不再硬编码，由 config.py 提供（见 create_daily 的 company 参数）。


def cn_date(y, m, d):
    return f'{y}年 {m} 月{d}日'


def daily_filename(dtype, y, m, d):
    """如 1.安保2026-07-01.docx"""
    return f'{int(d)}.{dtype}{y:04d}-{m:02d}-{d:02d}.docx'


def find_date_in_table(table):
    """在表格首行查找 yyyy年m月d日"""
    for row in table.rows[:2]:
        for cell in row.cells:
            mt = DATE_RE.search(cell.text)
            if mt:
                return int(mt.group(1)), int(mt.group(2)), int(mt.group(3))
    return None


def _body_groups(doc):
    """把正文元素按“表格”分组，每组 = 一天（表格前的段落 + 表格本身）。
    返回 (groups, trailing)；trailing 为最后一个表格之后的零散段落。"""
    body = doc.element.body
    elems = [e for e in body if not e.tag.endswith('}sectPr')]
    groups, cur = [], []
    for e in elems:
        cur.append(e)
        if e.tag == W_NS + 'tbl':
            groups.append(cur)
            cur = []
    return groups, cur


def _strip_unused_images(doc):
    """删除正文中未引用的图片关系，避免拆分后每个文件携带全部图片。"""
    xml = doc.element.body.xml
    used = set(re.findall(r'r:(?:embed|link)="(rId\d+)"', xml))
    for rId, rel in list(doc.part.rels.items()):
        if rel.reltype == RT.IMAGE and rId not in used:
            doc.part.drop_rel(rId)


def _strip_leading_empty_paragraphs(doc):
    """删除文档开头（首个表格之前）的连续纯空段落，避免下载文件顶部多余换行。
    含图片/图形的段落（如 LOGO 段）即使文本为空也保留。"""
    P = W_NS + 'p'
    body = doc.element.body
    # 找到第一个表格元素
    first_tbl = None
    for e in body:
        if e.tag == W_NS + 'tbl':
            first_tbl = e
            break
    # 删除第一个表格之前、文本为空且不含 drawing 的段落
    to_remove = []
    for e in list(body):
        if e is first_tbl:
            break
        if e.tag == P and not e.text.strip() and '<w:drawing' not in e.xml:
            to_remove.append(e)
    for e in to_remove:
        body.remove(e)


def split_monthly_docx(src_path, dtype, out_dirs):
    """将月度台账拆分为每日 docx。
    src_path: 月度文档路径; dtype: 安保/消控; out_dirs: 输出目录列表（可多份）。
    返回 [(date_tuple, filename), ...]
    """
    base = Document(src_path)
    groups, _ = _body_groups(base)
    n = len(groups)
    results = []
    for i in range(n):
        doc = Document(src_path)
        gs, trailing = _body_groups(doc)
        body = doc.element.body
        # 删除其它天的元素及尾部零散段落
        for j, g in enumerate(gs):
            if j != i:
                for e in g:
                    body.remove(e)
        for e in trailing:
            body.remove(e)
        _strip_leading_empty_paragraphs(doc)
        _strip_unused_images(doc)
        # 从表格中取日期
        date = None
        if doc.tables:
            date = find_date_in_table(doc.tables[0])
        if not date:
            date = (0, 0, i + 1)
        y, m, d = date
        fname = daily_filename(dtype, y, m, d)
        saved = None
        for od in out_dirs:
            os.makedirs(od, exist_ok=True)
            path = os.path.join(od, fname)
            if saved is None:
                doc.save(path)
                saved = path
            else:
                shutil.copy2(saved, path)
        results.append(((y, m, d), fname))
    return results


# ------------------------- 每日台账解析 / 编辑 -------------------------

def _content_cell(table):
    """工作内容单元格（第2行合并区域）"""
    return table.rows[1].cells[-1]


def parse_daily(path):
    """解析每日台账 -> dict"""
    doc = Document(path)
    if not doc.tables:
        return None
    t = doc.tables[0]
    date = find_date_in_table(t)
    dept = t.rows[0].cells[1].text.strip()
    content = _content_cell(t).text.strip()
    sign = t.rows[-1].cells[-1].text.strip() if len(t.rows) >= 5 else ''
    images = _count_images(doc)
    return {
        'date': f'{date[0]:04d}-{date[1]:02d}-{date[2]:02d}' if date else '',
        'dept': dept,
        'content': content,
        'sign': sign,
        'images': images,
    }


def _iter_image_rids(doc):
    rids = []
    for blip in doc.element.body.iter(
            '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
        rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if rid:
            rids.append(rid)
    return rids


def _count_images(doc):
    """只统计“工作图片”照片框内的图片，忽略标题 LOGO 等。"""
    if not doc.tables:
        return 0
    frames = _find_pic_frames(doc.tables[0])
    if not frames:
        return 0
    n = 0
    for c in frames:
        if c._tc.xml.count('<w:drawing') > 0:
            n += 1
    return n


def get_image(path, index):
    """返回 (bytes, content_type) 或 None（仅取照片框内的图片）"""
    doc = Document(path)
    if not doc.tables:
        return None
    frames = _find_pic_frames(doc.tables[0])
    BLIP = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
    EMBED = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
    rids = []
    for c in frames:
        for blip in c._tc.iter(BLIP):
            rid = blip.get(EMBED)
            if rid:
                rids.append(rid)
    if index < 0 or index >= len(rids):
        return None
    part = doc.part.related_parts[rids[index]]
    return part.blob, part.content_type


def update_content(path, content, sign=None):
    """更新工作内容文字（及可选签名），保留其它单元格与图片。"""
    doc = Document(path)
    cell = _content_cell(doc.tables[0])
    # 记录原字体
    font_name, font_size = '宋体', Pt(12)
    for p in cell.paragraphs:
        for r in p.runs:
            if r.text.strip():
                if r.font.name:
                    font_name = r.font.name
                if r.font.size:
                    font_size = r.font.size
                break
    # 清空原段落（保留第一个）
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)
    lines = content.splitlines() or ['']
    _set_run(p0.add_run(lines[0]), font_name, font_size)
    for line in lines[1:]:
        p = cell.add_paragraph()
        _set_run(p.add_run(line), font_name, font_size)
    # 更新签名（第5行最后一格），仅当传入且文件含签名行
    if sign is not None and len(doc.tables[0].rows) >= 5:
        scell = doc.tables[0].rows[-1].cells[-1]
        for p in scell.paragraphs:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            if p is not scell.paragraphs[0]:
                p._element.getparent().remove(p._element)
        sname, ssize = font_name, font_size
        for p in scell.paragraphs:
            for r in p.runs:
                if r.text.strip():
                    if r.font.name:
                        sname = r.font.name
                    if r.font.size:
                        ssize = r.font.size
                    break
        _set_run(scell.paragraphs[0].add_run(sign), sname, ssize)
    doc.save(path)


def _set_run(run, name, size):
    run.font.name = name
    run.font.size = size
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(W_NS + 'rFonts')
    if rFonts is not None:
        rFonts.set(W_NS + 'eastAsia', name)


def _find_pic_frames(table):
    """定位“工作图片”区域的所有照片框单元格。
    支持三种结构：
      - 2×2 方阵：内容区内嵌一张 2×2 内联表（共4框）；
      - 上下两行各一合并单元格（与月度模板一致，共2框）；
      - 兼容旧文档：主表每排4框 × 2排（共8框）。
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    for row in table.rows:
        cells = list(dict.fromkeys(row.cells))
        if cells and cells[0].text.strip() == '工作图片':
            # 1) 内嵌 2×2 方阵
            content = cells[1] if len(cells) > 1 else None
            if content is not None:
                tbl_elm = content._tc.find(qn('w:tbl'))
                if tbl_elm is not None:
                    inner = Table(tbl_elm, table.part)
                    return [c for r in inner.rows for c in r.cells]
            # 2) 上下两行各一合并单元格：收集“工作图片”行及紧随其后的内容单元格
            frames = [c for c in cells[1:]]
            # 若存在下一行（合并的第2行），也收集其内容单元格
            ridx = None
            for i, r in enumerate(table.rows):
                if r is row:
                    ridx = i
                    break
            if ridx is None:
                ridx = next(
                    (i for i, r in enumerate(
                        table.rows) if list(
                        dict.fromkeys(
                            r.cells))[0].text.strip() == '工作图片'), None)
            if ridx is not None and ridx + 1 < len(table.rows):
                nxt = list(dict.fromkeys(table.rows[ridx + 1].cells))
                frames.extend(c for c in nxt[1:])
            frames = list(dict.fromkeys(frames))
            if frames:
                return frames
    # 兼容旧文档：主表每排4框
    frames = []
    for row in table.rows:
        cells = list(dict.fromkeys(row.cells))
        if cells and cells[0].text.strip() == '工作图片':
            frames.extend(cells[1:])
    return frames


def _place_one(cell, data, avail_w, avail_h):
    """在指定照片框单元格中放置一张图片（先清空原图片与占位文字）。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _AP
    # 清空单元格内原有段落（含占位文字/旧图片）
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    p = cell.add_paragraph()
    p.alignment = _AP.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 读取照片真实像素尺寸以计算缩放比例
    try:
        im = PILImage.open(io.BytesIO(data))
        iw, ih = im.size
    except Exception:
        iw = ih = 1000
    w_in, h_in = iw / 96.0, ih / 96.0
    if w_in <= 0 or h_in <= 0:
        w_in = h_in = 1.0
    scale = min(avail_w / w_in, avail_h / h_in)
    tw = w_in * scale
    p.add_run().add_picture(io.BytesIO(data), width=Inches(tw))


def add_images(path, files, index=None):
    """向工作图片框添加图片，按每张照片真实长宽比等比缩放后填入。

    files: [(bytes, filename), ...]
    index: 若指定（0 起），则把图片放入指定序号的框（替换该框原内容）；
           否则按原逻辑顺序填入空框。返回实际写入的张数。"""
    doc = Document(path)
    t = doc.tables[0]
    frames = _find_pic_frames(t)
    if not frames:
        raise ValueError('未找到工作图片框')
    # 根据文档实际布局自适应可用区域（英寸）：
    #   - 原 8 框布局（每排4框×2排）：每框较大，用 3.0×2.6 恢复原效果
    #   - 新 2×2 方阵（共4框）：每框较小，用 2.4×2.1
    if len(frames) >= 8:
        avail_w, avail_h = 3.0, 2.6
    else:
        avail_w, avail_h = 2.4, 2.1

    # 指定框：仅填入目标框
    if index is not None:
        if index < 0 or index >= len(frames):
            raise ValueError(f'照片框序号超出范围（共 {len(frames)} 个框）')
        _place_one(frames[index], files[0][0], avail_w, avail_h)
        doc.save(path)
        return 1

    # 顺序填空框
    empty = [c for c in frames
             if c._tc.xml.count('<w:drawing') == 0]
    used = 0
    for data, _name in files:
        if not empty:
            break
        cell = empty.pop(0)
        _place_one(cell, data, avail_w, avail_h)
        used += 1
    doc.save(path)
    return used


# ------------------------- 新建每日台账 -------------------------

def create_daily(path, dtype, y, m, d, content='', company='', logo_path=None):
    """按模板结构新建每日台账 docx（工作图片区为 2×2 方阵 = 4个照片框）。

    company: 公司全称（如“新东升物业（嘉应学院丰顺校区）”），由系统设置提供，不再硬编码。
    logo_path: 可选，模板标题上方嵌入的 LOGO 图片路径。
    """
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.0)
    sec.top_margin, sec.bottom_margin = Cm(1.8), Cm(1.8)

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(W_NS + 'eastAsia', '宋体')

    # 可选 LOGO（标题上方居中）
    if logo_path and os.path.exists(logo_path):
        try:
            lp = doc.add_paragraph()
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lp.add_run().add_picture(logo_path, height=Cm(2.2))
        except Exception:
            pass

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run('工作台账')
    r.bold = True
    r.font.size = Pt(22)
    _set_run(r, '宋体', Pt(22))
    r.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(company or ' ')
    _set_run(r2, '宋体', Pt(14))
    r2.bold = True

    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    # 固定表格宽度，使其撑满页面（与月度模板一致），避免自动宽度导致表单偏窄
    table.autofit = False
    table.allow_autofit = False
    tblPr = table._tbl.tblPr
    from docx.oxml.ns import qn
    tw = tblPr.find(qn('w:tblW'))
    if tw is None:
        tw = tblPr.makeelement(qn('w:tblW'), {})
        tblPr.append(tw)
    tw.set(qn('w:type'), 'dxa')
    tw.set(qn('w:w'), '10558')

    def put(cell, text, size=Pt(12), bold=False, center=True, left=False):
        p = cell.paragraphs[0]
        if left:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(text)
        _set_run(rr, '宋体', size)
        rr.bold = bold
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 第0行：部门 / 日期（日期标签跨两列）
    put(table.cell(0, 0), '部 门', bold=True)
    put(table.cell(0, 1), DEPT_NAME.get(dtype, dtype))
    dc = table.cell(0, 2).merge(table.cell(0, 3))
    put(dc, '日 期', bold=True)
    put(table.cell(0, 4), cn_date(y, m, d))

    # 第1行：第一列“工作内容”标签 + 右侧合并单元格填入内容
    put(table.cell(1, 0), '工作内容', bold=True)
    wc = table.cell(1, 1).merge(table.cell(1, 4))
    lines = content.splitlines() or ['']
    p0 = wc.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run(p0.add_run(lines[0]), '宋体', Pt(12))
    for line in lines[1:]:
        _set_run(wc.add_paragraph().add_run(line), '宋体', Pt(12))

    # 第2、3行：工作图片（标签竖向合并跨两行；内容区内嵌 2×2 方阵 = 4 个照片框）
    label = table.cell(2, 0).merge(table.cell(3, 0))
    put(label, '工作图片', bold=True)
    content_cell = table.cell(2, 1).merge(table.cell(3, 4))
    content_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    content_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 内嵌 2×2 表作为 4 个照片框
    inner = content_cell.add_table(rows=2, cols=2)
    inner.style = 'Table Grid'
    inner.autofit = False
    inner.allow_autofit = False
    for r in inner.rows:
        for c in r.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 让内嵌表宽度撑满内容单元格
    itbl = inner._tbl.tblPr
    from docx.oxml.ns import qn as _qn
    iw = itbl.find(_qn('w:tblW'))
    if iw is None:
        iw = itbl.makeelement(_qn('w:tblW'), {})
        itbl.append(iw)
    iw.set(_qn('w:type'), 'dxa')
    iw.set(_qn('w:w'), '9000')

    # 第4行：第一列“签 名”标签 + 右侧签名行
    put(table.cell(4, 0), '签 名', bold=True)
    sg = table.cell(4, 1).merge(table.cell(4, 4))
    put(sg, '班 长：                   主 管：')

    # 列宽（对齐月度模板：首列约 2.41cm，后四列均衡撑满约 18.7cm 内容区）
    colw = [Cm(2.41), Cm(4.06), Cm(4.06), Cm(4.06), Cm(4.06)]
    for ci, w in enumerate(colw):
        for row in table.rows:
            row.cells[ci].width = w

    # 行高：设下限随内容自然撑开（对齐模板观感）
    table.rows[0].height = Cm(1.39)
    table.rows[1].height = Cm(3.69)
    table.rows[2].height = Cm(6.98)
    table.rows[3].height = Cm(7.11)
    table.rows[4].height = Cm(1.82)

    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)


def count_tables(path):
    return len(Document(path).tables)
